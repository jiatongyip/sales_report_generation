import os
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from typing import Any, Dict, Callable, List, Optional, Tuple
import json
import statsmodels.formula.api as smf 
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.utils import load_dataframe
import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
logger = logging.getLogger(__name__)


def top_performers(input_path: str, group_by_column: str, metric_column: str, n: int = 5) -> str:
    """
    Identifies and returns the top N performers from the DataFrame based on a specified metric.

    Args:
        input_path (str): Input file path.
        group_by_column (str): The column to group by (e.g., 'salesperson', 'region', 'product').
        metric_column (str): The column containing the metric to evaluate performance (e.g., 'sales_amount', 'profit').
        n (int): The number of top performers to return.

    Returns:
        str: A JSON string representing the top N performers, sorted by the metric.
    """
    df = load_dataframe(input_path)

    if df.empty or group_by_column not in df.columns or metric_column not in df.columns:
        return json.dumps({"error": "Invalid DataFrame or columns provided."})

    # Group by the specified column and sum the metric
    performance = df.groupby(group_by_column)[metric_column].sum().nlargest(n).reset_index()

    # Convert to list of dictionaries for JSON serialization
    result = performance.to_dict(orient='records')
    return json.dumps(result)

def under_performers(input_path: str, group_by_column: str, metric_column: str, n: int = 5) -> str:
    """
    Identifies and returns the bottom N under-performers from the DataFrame based on a specified metric.

    Args:
        input_path (str): Input file path.
        group_by_column (str): The column to group by (e.g., 'salesperson', 'region', 'product').
        metric_column (str): The column containing the metric to evaluate performance (e.g., 'sales_amount', 'profit').
        n (int): The number of under-performers to return.

    Returns:
        str: A JSON string representing the bottom N under-performers, sorted by the metric.
    """
    df = load_dataframe(input_path)

    if df.empty or group_by_column not in df.columns or metric_column not in df.columns:
        return json.dumps({"error": "Invalid DataFrame or columns provided."})

    # Group by the specified column and sum the metric
    # Using nsmallest to get under-performers
    performance = df.groupby(group_by_column)[metric_column].sum().nsmallest(n).reset_index()

    # Convert to list of dictionaries for JSON serialization
    result = performance.to_dict(orient='records')
    return json.dumps(result)

def ols_key_driver_analysis(input_path: str, target_column: str, feature_columns: list[str]):
    """
    **TOOL DESCRIPTION**
    Performs an Ordinary Least Squares (OLS) regression analysis using the Statsmodels formula API 
    to identify key drivers (features) impacting a target variable. It automatically handles both 
    numeric and non-numeric (categorical) feature columns via one-hot encoding.

    **INPUTS**
    - `input_path` (str): Input file path.
    - `target_column` (str): The name of the dependent variable (Y) column in the DataFrame.
    - `feature_columns` (list[str]): A list of independent variable (X) column names used as predictors.
    - `input_path` (str): The input data path.

    **OUTPUT**
    Returns a JSON string containing the regression results: 
    *   `coefficients`: A dictionary of beta coefficients for each predictor (including one-hot encoded dummies).
    *   `p_values`: A dictionary of p-values associated with each coefficient's significance.
    *   `r_squared`: The R-squared value for the overall model fit.
    
    Returns a JSON string with an 'error' key if the analysis fails (e.g., insufficient data, perfect multicollinearity).

    **BEHAVIOR DETAILS**
    *   The `target_column` is converted to numeric, with errors coerced to NaN, then dropped.
    *   Rows with NaNs in any of the `feature_columns` are dropped.
    *   `statsmodels.formula.api.ols` automatically handles one-hot encoding for string/object columns specified in `feature_columns`.

    **EXAMPLE OUTPUT (Successful):**
    ```json
    {
      "coefficients": {"Intercept": 0.5, "feature_a": 1.2, "feature_b[T.Category2]": -0.4},
      "p_values": {"Intercept": 0.001, "feature_a": 0.05, "feature_b[T.Category2]": 0.25},
      "r_squared": 0.75
    }
    ```
    """
    df = load_dataframe(input_path)
    # Ensure target and feature columns are numeric, coercing errors will turn non-numeric into NaN
    # Ensure the target column is numeric
    df[target_column] = pd.to_numeric(df[target_column], errors='coerce')
    
    # Drop rows with NaN values in the target column
    df_clean = df.dropna(subset=[target_column])
    
    # Optional: Drop rows where *any* feature column is NaN, as OLS cannot handle missing features
    df_clean = df_clean.dropna(subset=feature_columns)
    
    if df_clean.empty:
        return json.dumps({"error": "No valid data after cleaning for OLS analysis."})
    
    # Construct the OLS formula string dynamically
    formula = f'{target_column} ~ {" + ".join(feature_columns)}'
    
    try:
        model = smf.ols(formula=formula, data=df_clean)
        results = model.fit()
        
        # Extract beta coefficients and p-values
        ols_results = {
            "coefficients": results.params.to_dict(),
            "p_values": results.pvalues.to_dict(),
            "r_squared": results.rsquared
        }
        return json.dumps(ols_results)
    except Exception as e:
        return json.dumps({"error": str(e)})

def preview_data(input_path: str, head_rows: int = 5) -> str:
    """
    Returns a preview of the DataFrame, typically the first few rows.

    Args:
        input_path (str): Input file path.
        head_rows (int): The number of rows to return from the head of the DataFrame.

    Returns:
        str: A JSON string representing the head of the DataFrame.
    """
    df = load_dataframe(input_path)
    if df.empty:
        return json.dumps({"message": "DataFrame is empty."})

    # Convert the head of the DataFrame to a list of dictionaries for JSON serialization
    # Ensure datetime objects are converted to strings (ISO format)
    df_head = df.head(head_rows)
    for col in df_head.columns:
        if pd.api.types.is_datetime64_any_dtype(df_head[col]):
            df_head[col] = df_head[col].dt.isoformat()

    result = df_head.to_dict(orient='records')
    return json.dumps(result)


def generate_simple_plot(
    input_path: str,
    plot_type: str,
    title: str,
    x_col: str,
    y_col: str,
    groupby_col: Optional[str] = None,
    save_path: str = "dynamic_plot.png",
    agg_func: str = "sum",
) -> str:
    """
    Generates a simple bar or line plot from a dataset with optional aggregation.
    
    Line plots support optional grouping; bar plots do not.
    Aggregation supports 'sum' or 'mean'.

    Parameters
    ----------
    input_path : str
        Path to CSV/Excel file.
    plot_type : str
        'line' or 'bar'.
    title : str
        Plot title.
    x_col : str
        Column for x-axis.
    y_col : str
        Column for y-axis values.
    groupby_col : Optional[str]
        Column to group by for line plots (ignored for bar plots).
    save_path : str
        File path to save the plot.
    agg_func : str
        Aggregation function: 'sum' or 'mean'.

    Returns
    -------
    str
        JSON string containing:
        - 'plot_path': absolute path to saved plot
        - 'plot_data_json': JSON string of aggregated data
    """
    # 1. Load data
    df = load_dataframe(input_path)
    df[y_col] = pd.to_numeric(df[y_col], errors="coerce")
    cols_to_check = [x_col, y_col]
    if plot_type == "line" and groupby_col:
        cols_to_check.append(groupby_col)
    df = df.dropna(subset=cols_to_check)
    if df.empty:
        return json.dumps({"error": "No valid data after cleaning."})

    # 2. Validate aggregation
    agg_func = (agg_func or "sum").lower()
    if agg_func not in ("sum", "mean"):
        agg_func = "sum"

    # 3. Aggregate data
    if plot_type == "line":
        group_cols = [x_col] + ([groupby_col] if groupby_col else [])
    else:
        group_cols = [x_col]

    df_plot = df.groupby(group_cols)[y_col].agg(agg_func).reset_index()
    df_plot = df_plot.sort_values(by=x_col)

    # 4. Plot
    plt.figure(figsize=(10, 6))
    if plot_type == "line":
        if groupby_col:
            for label, g in df_plot.groupby(groupby_col):
                plt.plot(g[x_col], g[y_col], marker="o", label=str(label))
            if df_plot[groupby_col].nunique() > 1:
                plt.legend(title=groupby_col.title(), bbox_to_anchor=(1.05, 1), loc="upper left")
        else:
            plt.plot(df_plot[x_col], df_plot[y_col], marker="o")
    elif plot_type == "bar":
        df_plot.plot(kind="bar", x=x_col, y=y_col, ax=plt.gca(), legend=False)
        plt.xticks(rotation=45, ha="right")
    else:
        return json.dumps({"error": f"Unsupported plot type '{plot_type}'"})

    # 5. Customize plot
    plt.title(title)
    plt.xlabel(x_col.title())
    plt.ylabel(f"{y_col.title()} ({agg_func})")
    plt.grid(True, axis="y")
    if pd.api.types.is_datetime64_any_dtype(df_plot[x_col]):
        plt.gcf().autofmt_xdate()
    plt.tight_layout()

    # 6. Save plot
    os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
    plt.savefig(save_path, dpi=300)
    plt.close()
    absolute_plot_path = os.path.abspath(save_path)

    # 7. Prepare JSON data
    df_out = df_plot.copy()
    for col in df_out.columns:
        if pd.api.types.is_datetime64_any_dtype(df_out[col]):
            df_out[col] = df_out[col].dt.isoformat()
    plot_data_json = json.dumps(df_out.to_dict(orient="records"))

    return json.dumps({"plot_path": absolute_plot_path, "plot_data_json": plot_data_json})

def assemble_simple_report_docx(sections_json: str, report_title: str = 'Generated Report', outfile: str = 'report.docx') -> str:
    """
    **TOOL DESCRIPTION**
    Generates a Microsoft Word (.docx) report from structured input provided as a JSON list of sections. 
    Each section can include a title, narrative text content, and one or more local paths to image plot files.

    **INPUTS**
    - `sections_json` (str): A JSON string representing a list of section dictionaries. 
                             Each dictionary must have a 'title' (string) key.
                             It can optionally include 'content' (string for narrative text) and/or 
                             'plot_paths' (a single string or a list of strings for local image file paths, 
                             e.g., 'plots/my_plot.png' or ["plot1.png", "plot2.png"]).
    - `report_title` (str, optional): The main title for the DOCX report. Defaults to 'Generated Report'.
    - `outfile` (str, optional): The local file path where the generated DOCX report should be saved 
                                 (e.g., 'sales_report.docx'). Defaults to 'report.docx'.

    **OUTPUT**
    Returns the absolute path to the saved DOCX report file as a string. Returns an error message string if processing fails.

    **EXAMPLES FOR LLM USAGE (sections_json parameter):**

    EXAMPLE 1: A simple report with text only.
    ```json
    [
      {"title": "Executive Summary", "content": "Sales performance improved significantly in Q3, driven by new product launches in the APAC region."}
    ]
    ```

    EXAMPLE 2: A section with a single plot.
    ```json
    [
      {"title": "Sales Trend Chart", "content": "The following chart details the total sales.", "plot_paths": "/app/data/plots/monthly_sales_line_plot.png"}
    ]
    ```

    EXAMPLE 3: A section with multiple plots (vertical stacking).
    ```json
    [
      {"title": "Regional Performance with Visualizations", "content": "The following plots display performance metrics side-by-side (vertically stacked).", "plot_paths": ["./output_plots/regional_bar_chart.png", "./output_plots/regional_line_chart.png"]}
    ]
    ```
    """
    
    try:
        sections: List[Dict[str, Any]] = json.loads(sections_json)
    except json.JSONDecodeError:
        return f"Error: Invalid JSON data provided for sections_json. Expected a JSON list of objects."

    document = Document()
    document.add_heading(report_title, level=0)
    date_p = document.add_paragraph(f"Generated: {datetime.now().isoformat()}")
    date_p.style.font.size = Pt(10)

    for i, sec in enumerate(sections):
        title = sec.get('title', f'Section {i+1}')
        content = sec.get('content', '')
        # Retrieve the value for 'plot_paths'
        plot_paths_input = sec.get('plot_paths', []) 
        
        # Ensure plot_paths_input is always a list for consistent iteration
        if isinstance(plot_paths_input, str):
            plot_paths_list = [plot_paths_input]
        elif isinstance(plot_paths_input, list):
            plot_paths_list = plot_paths_input
        else:
            plot_paths_list = []
            
        document.add_heading(title, level=1)
        
        if content:
            document.add_paragraph(content)
            
        if plot_paths_list:
            for plot_path in plot_paths_list:
                if os.path.exists(plot_path):
                    # Each plot gets its own centered paragraph
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    # Add image, constraining width to 6 inches
                    run.add_picture(plot_path, width=Inches(6)) 
                else:
                    document.add_paragraph(f"[Error: Plot file not found at {plot_path}]")
        
        document.add_paragraph("-" * 20) 

    os.makedirs(os.path.dirname(outfile) or '.', exist_ok=True)
    
    try:
        document.save(outfile)
        return os.path.abspath(outfile)
    except Exception as e:
        return f"An error occurred during DOCX generation: {str(e)}"


